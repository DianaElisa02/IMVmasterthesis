"""
make_descriptive_table.py
=========================
Builds Table 2 (descriptive statistics) for the Data section of the thesis
and writes it as a Word .docx file.

Structure
---------
Single pooled column across all analysis years (2017-2019, 2021-2025).
Statistics per variable: Mean, SD, p25, p75.

Variables, grouped:
  A. Outcome variables
       vhPobreza      (primary outcome — at-risk-of-poverty)
       vhMATDEP       (severe material deprivation)
       poverty_gap    (FGT(1), normalised)
       poverty_gap_sq (FGT(2))
  B. Treatment / exposure variables
       EUROMOD-based simulated exposure index
       Administrative RMI expenditure per capita
  C. Core control
       hh_size        (household size)

Statistics are survey-weighted using weight_hh (household-level), consistent
with the household unit of observation. Mean, SD, and the 25th/75th
percentiles are all survey-weighted.

Usage
-----
  python make_descriptive_table.py

Requires the enriched panel produced by run_poverty_gap.py:
  output/analysis_dataset_with_gap.parquet

Output
------
  output/descriptive_table.docx

Dependencies
------------
  pip install polars numpy python-docx
"""

from __future__ import annotations

import logging
from pathlib import Path

import numpy as np
import polars as pl

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
)
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────────────
# Paths — anchored to this file's location (project root)
# ─────────────────────────────────────────────────────────────────────────────
_THIS_FILE   = Path(__file__).resolve()
BASE_PATH    = _THIS_FILE.parent
INPUT_PATH   = BASE_PATH / "output" / "analysis_dataset_with_gap.parquet"
OUTPUT_DOCX  = BASE_PATH / "output" / "descriptive_table.docx"

# ─────────────────────────────────────────────────────────────────────────────
# CONFIG — set the exact exposure column names from YOUR panel here.
# Check them with:  pl.read_parquet(INPUT_PATH).columns
# The EUROMOD-based index is your primary exposure spec (EXPOSURE_SPECS[0]).
# ─────────────────────────────────────────────────────────────────────────────
EXPOSURE_EUROMOD = "exposure_composite_hybrid"     # <-- EDIT if different
EXPOSURE_RMI_PC  = "exposure_admin"               # <-- EDIT if different

WEIGHT_COL = "weight_hh"

# (column candidates, display label, group). The first column found is used,
# so both raw ECV names and renamed versions are handled automatically.
VARIABLE_SPEC = [
    # ── A. Outcome variables ──────────────────────────────────────────────────
    (["vhPobreza", "poverty"],        "At-risk-of-poverty (vhPobreza)",        "A. Outcome variables"),
    (["vhMATDEP", "matdep"],          "Severe material deprivation (vhMATDEP)","A. Outcome variables"),
    (["poverty_gap"],                 "Poverty gap (FGT1)",                    "A. Outcome variables"),
    (["poverty_gap_sq"],              "Poverty gap squared (FGT2)",            "A. Outcome variables"),
    # ── B. Treatment / exposure variables ─────────────────────────────────────
    ([EXPOSURE_EUROMOD],              "Exposure index (EUROMOD-simulated)",    "B. Treatment / exposure"),
    ([EXPOSURE_RMI_PC],               "RMI expenditure per capita (admin.)",   "B. Treatment / exposure"),
    # ── C. Core control ────────────────────────────────────────────────────────
    (["hh_size"],                     "Household size",                        "C. Control"),
]


def _resolve_column(candidates: list[str], available: list[str]) -> str | None:
    """Return the first candidate present in the panel, else None."""
    for c in candidates:
        if c in available:
            return c
    return None


def _weighted_quantile(
    values: np.ndarray, weights: np.ndarray, q: float
) -> float:
    """
    Weighted quantile q (0 < q < 1) using the cumulative-weight method.

    Values are sorted, cumulative weights are normalised to [0, 1] using the
    midpoint convention ((cumw - 0.5*w) / total), and the quantile is found by
    linear interpolation. Equivalent to the standard weighted-percentile
    definition and consistent with how the poverty line median is computed.
    """
    if values.size == 0:
        return np.nan
    sort_idx = np.argsort(values)
    v = values[sort_idx]
    w = weights[sort_idx]
    cumw = np.cumsum(w)
    total = cumw[-1]
    if total <= 0:
        return np.nan
    # Midpoint (Type 5 / Hazen) plotting positions
    p = (cumw - 0.5 * w) / total
    return float(np.interp(q, p, v))


def _weighted_stats(
    values: np.ndarray, weights: np.ndarray
) -> tuple[float, float, float, float, int]:
    """
    Weighted mean, SD, p25, p75; plus unweighted N of valid obs.

    SD is the weighted sample standard deviation using reliability weights:
      var = sum(w * (x - xbar)^2) / (sum(w) - sum(w^2)/sum(w))
    which reduces to the usual unbiased estimator when all weights equal 1.
    Percentiles are survey-weighted (see _weighted_quantile).

    Returns
    -------
    (mean, sd, p25, p75, n)
    """
    mask = ~np.isnan(values) & ~np.isnan(weights) & (weights > 0)
    v, w = values[mask], weights[mask]
    n = int(v.size)
    if n == 0:
        return (np.nan, np.nan, np.nan, np.nan, 0)

    sw    = w.sum()
    mean  = float(np.sum(w * v) / sw)
    denom = sw - (np.sum(w ** 2) / sw)
    if denom > 0:
        var = float(np.sum(w * (v - mean) ** 2) / denom)
        sd  = float(np.sqrt(var))
    else:
        sd = np.nan

    p25 = _weighted_quantile(v, w, 0.25)
    p75 = _weighted_quantile(v, w, 0.75)
    return (mean, sd, p25, p75, n)


def _fmt(x: float, decimals: int = 3) -> str:
    if x is None or (isinstance(x, float) and np.isnan(x)):
        return "—"
    # Large values (e.g. RMI per capita in euros) shown with thousands sep, 0 dp
    if abs(x) >= 1000:
        return f"{x:,.0f}"
    return f"{x:.{decimals}f}"


def compute_table() -> tuple[list[dict], int]:
    """Compute the descriptive statistics rows. Returns (rows, n_obs_panel)."""
    logger.info("Reading enriched panel: %s", INPUT_PATH)
    if not INPUT_PATH.exists():
        raise FileNotFoundError(
            f"Cannot find {INPUT_PATH}. Run run_poverty_gap.py first to "
            f"produce analysis_dataset_with_gap.parquet."
        )

    panel = pl.read_parquet(INPUT_PATH)
    available = panel.columns
    n_panel = len(panel)
    logger.info("Panel loaded: %d obs | %d columns", n_panel, len(available))

    if WEIGHT_COL not in available:
        raise ValueError(f"Weight column '{WEIGHT_COL}' not found in panel.")
    weights_all = panel[WEIGHT_COL].to_numpy().astype(float)

    rows: list[dict] = []
    for candidates, label, group in VARIABLE_SPEC:
        col = _resolve_column(candidates, available)
        if col is None:
            logger.warning(
                "Variable not found (candidates=%s) — row left blank. "
                "Check column names in the panel.", candidates
            )
            rows.append({
                "group": group, "label": label,
                "mean": np.nan, "sd": np.nan,
                "p25": np.nan, "p75": np.nan, "n": 0,
                "found": False, "col": None,
            })
            continue

        vals = panel[col].to_numpy().astype(float)
        mean, sd, p25, p75, n = _weighted_stats(vals, weights_all)
        rows.append({
            "group": group, "label": label,
            "mean": mean, "sd": sd, "p25": p25, "p75": p75, "n": n,
            "found": True, "col": col,
        })
        logger.info(
            "%-38s [%s] mean=%.4f sd=%.4f p25=%.4f p75=%.4f n=%d",
            label, col, mean, sd, p25, p75, n,
        )

    return rows, n_panel


def write_docx(rows: list[dict], n_panel: int) -> None:
    """Write the descriptive table to a Word .docx file using python-docx."""
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    # Title
    title = doc.add_paragraph()
    run = title.add_run("Table 2. Descriptive Statistics")
    run.bold = True
    run.font.size = Pt(12)

    sub = doc.add_paragraph()
    sub_run = sub.add_run(
        "Pooled across analysis years (2017–2019, 2021–2025). "
        "Means, standard deviations, and percentiles are survey-weighted "
        "(household weights)."
    )
    sub_run.italic = True
    sub_run.font.size = Pt(9)

    # Table: Variable | Mean | SD | p25 | p75
    n_cols = 5
    table = doc.add_table(rows=1, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    headers = ["Variable", "Mean", "SD", "p25", "p75"]
    hdr = table.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ""
        p = hdr[j].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        if j == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    current_group = None
    for row in rows:
        # Group header row spanning all columns
        if row["group"] != current_group:
            current_group = row["group"]
            grp_cells = table.add_row().cells
            merged = grp_cells[0]
            for j in range(1, n_cols):
                merged = merged.merge(grp_cells[j])
            p = merged.paragraphs[0]
            gr = p.add_run(current_group)
            gr.bold = True
            gr.italic = True

        cells = table.add_row().cells
        # Variable label
        cells[0].text = ""
        cells[0].paragraphs[0].add_run("  " + row["label"])
        # Stats
        vals = [
            _fmt(row["mean"]),
            _fmt(row["sd"]),
            _fmt(row["p25"]),
            _fmt(row["p75"]),
        ]
        for j, v in enumerate(vals, start=1):
            cells[j].text = ""
            p = cells[j].paragraphs[0]
            p.add_run(v)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Notes
    note = doc.add_paragraph()
    note_run = note.add_run(
        "Notes: vhPobreza and vhMATDEP are binary indicators; their means are "
        "weighted prevalence rates. Poverty gap (FGT1) and its square (FGT2) "
        "are computed relative to a year-specific poverty line set at 40% of the "
        "person-weighted median equivalised income. The EUROMOD-based exposure "
        "index is the simulated gain in minimum income protection (IMV minus "
        "pre-reform RMI entitlement); RMI expenditure per capita is from "
        "administrative reports. Household size is the only model covariate. "
        "Source: ECV (INE) microdata and Ministerio de Derechos Sociales RMI reports."
    )
    note_run.font.size = Pt(8)
    note_run.italic = True

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    logger.info("Saved Word table: %s", OUTPUT_DOCX)


def main() -> None:
    logger.info("=== make_descriptive_table.py ===")
    rows, n_panel = compute_table()

    missing = [r["label"] for r in rows if not r["found"]]
    if missing:
        logger.warning(
            "These variables were NOT found and appear blank in the table: %s. "
            "Edit EXPOSURE_EUROMOD / EXPOSURE_RMI_PC or VARIABLE_SPEC at the top "
            "of this script to match your panel's column names.", missing
        )

    write_docx(rows, n_panel)
    logger.info("Done. Open: %s", OUTPUT_DOCX)


if __name__ == "__main__":
    main()